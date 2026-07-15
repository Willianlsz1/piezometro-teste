# -*- coding: utf-8 -*-
"""
Gera docs/tcc/TCC_ENTREGA.docx a partir de docs/tcc/TCC_AQUASENSE.md,
usando o TEMPLATE INTEGRA - MG.docx como base (capa e estilos oficiais).

Uso:  python docs/tcc/gerar_tcc_integra.py [caminho-do-template]

O template oficial nao fica no repo (arquivo do SENAI). Caminho padrao:
%USERPROFILE%\\Downloads\\TEMPLATE INTEGRA - MG.docx
"""
import io, os, re, sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

AQUI = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(AQUI, "TCC_AQUASENSE.md")
SAIDA = os.path.join(AQUI, "TCC_ENTREGA.docx")
TEMPLATE = sys.argv[1] if len(sys.argv) > 1 else os.path.join(
    os.environ["USERPROFILE"], "Downloads", "TEMPLATE INTEGRA - MG.docx")

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
VERMELHO = RGBColor(0xC0, 0x00, 0x00)

# ---------------------------------------------------------------- parse do md
def carregar_secoes():
    """Divide o markdown em secoes por heading (##/###) e retorna lista de blocos."""
    linhas = open(MD, encoding="utf-8").read().splitlines()
    if not linhas or not linhas[0].startswith("# "):
        sys.exit(
            "ERRO: a primeira linha do TCC_AQUASENSE.md precisa comecar com '# ' (cerquilha e\n"
            "espaco), que e a marcacao de titulo do markdown. Ela deve ficar assim:\n\n"
            "  # #12900# AquaSense: Sistema de Telemetria e Monitoramento...\n\n"
            "O primeiro '# ' e a marcacao; o codigo da unidade vem entre as duas cerquilhas\n"
            "seguintes, conforme o manual do INTEGRA. Linha 1 atual:\n\n  %s"
            % (linhas[0] if linhas else "(arquivo vazio)")
        )
    secoes, atual = [], None
    for l in linhas:
        h = re.match(r"^(#{1,3}) +(.*)$", l)
        if h:
            atual = {"nivel": len(h.group(1)), "titulo": h.group(2).strip(), "linhas": []}
            secoes.append(atual)
        elif atual is not None:
            atual["linhas"].append(l)
    return secoes

def blocos(linhas):
    """('p'|'bullet'|'sub'|'grafico'|'table', dado)"""
    out, tabela, par = [], [], []
    def flush_par():
        if par:
            out.append(("p", " ".join(par).strip())); par.clear()
    def flush_tab():
        if tabela:
            out.append(("table", tabela[:])); tabela.clear()
    for l in linhas:
        s = l.strip()
        g = re.match(r"^\[GRAFICO: *([^ |]+) *\| *(.*)\]$", s)
        if s.startswith("|"):
            flush_par()
            cels = [c.strip() for c in s.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cels):
                continue
            tabela.append(cels)
        elif g:
            flush_par(); flush_tab()
            out.append(("grafico", (g.group(1), g.group(2))))
        elif s.startswith("- "):
            flush_par(); flush_tab()
            out.append(("bullet", s[2:].strip()))
        elif re.fullmatch(r"\*\*[^*]+\*\*", s):
            flush_par(); flush_tab()
            out.append(("sub", s.strip("*")))
        elif s in ("", "---") or s.startswith(">") or s.startswith("*Nota:"):
            flush_par(); flush_tab()
        else:
            flush_tab(); par.append(s)
    flush_par(); flush_tab()
    return out

# ------------------------------------------------------------- helpers do docx
def runs_formatados(p, texto):
    """Escreve texto em p tratando **bold**, `[PREENCHER: ...]` (vermelho) e limpando markdown."""
    for parte in re.split(r"(\*\*[^*]+\*\*|`\[PREENCHER:[^`]+\]`)", texto):
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**"):
            p.add_run(parte.strip("*")).bold = True
        elif parte.startswith("`[PREENCHER:"):
            r = p.add_run(parte.strip("`"))
            r.font.color.rgb = VERMELHO
            r.bold = True
        else:
            p.add_run(parte.replace("`", ""))

def novo_par(d, texto="", estilo="Normal", justificar=True):
    p = d.add_paragraph(style=estilo)
    if texto:
        runs_formatados(p, texto)
    if justificar and estilo == "Normal":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(6)
    return p

def sombrear(celula, hexcor):
    tcPr = celula._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(W + "shd", {W + "val": "clear", W + "fill": hexcor})
    tcPr.append(shd)

def add_tabela(d, dados):
    t = d.add_table(rows=len(dados), cols=len(dados[0]))
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, linha in enumerate(dados):
        for j, cel in enumerate(linha):
            c = t.rows[i].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            runs_formatados(p, cel)
            for r in p.runs:
                r.font.size = Pt(9)
                if i == 0:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                sombrear(c, "1F3864")
    return t

def add_grafico(d, arquivo, legenda):
    cam = os.path.join(AQUI, "graficos", arquivo)
    if not os.path.exists(cam):
        print("  ! grafico ausente:", arquivo)
        return
    pl = d.add_paragraph()
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runs_formatados(pl, legenda)
    for r in pl.runs:
        r.font.size = Pt(10)
        r.bold = True
    pi = d.add_paragraph()
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.add_run().add_picture(cam, width=Cm(14))
    pf = d.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("Fonte: elaborado pelos autores (2026).")
    rf.font.size = Pt(9)

def escreve(d, secao, ancora, placeholder=None):
    """Insere os blocos da secao logo apos a ancora."""
    ref = ancora._p
    primeiro = True
    for tipo, dado in blocos(secao["linhas"]):
        if tipo == "table":
            t = add_tabela(d, dado)
            ref.addnext(t._tbl); ref = t._tbl
            v = d.add_paragraph(); ref.addnext(v._p); ref = v._p
            continue
        if tipo == "grafico":
            marca = len(d.paragraphs)
            add_grafico(d, dado[0], dado[1])
            for p in d.paragraphs[marca:]:
                ref.addnext(p._p); ref = p._p
            continue
        if tipo == "sub":
            p = d.add_paragraph()
            r = p.add_run(dado); r.bold = True
            p.paragraph_format.space_before = Pt(10)
        elif tipo == "bullet":
            p = d.add_paragraph()
            runs_formatados(p, dado)
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                if r.text and not r.text.startswith("["):
                    break
            p.insert_paragraph_before  # noop
            p.runs[0].text = "•  " + p.runs[0].text if p.runs else ""
        else:
            if primeiro and placeholder is not None:
                placeholder.text = ""
                runs_formatados(placeholder, dado)
                placeholder.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                placeholder.paragraph_format.first_line_indent = Cm(1.25)
                ref = placeholder._p
                primeiro = False
                continue
            p = novo_par(d, dado)
        ref.addnext(p._p); ref = p._p
        primeiro = False

def acha(d, texto):
    for p in d.paragraphs:
        if p.text.strip() == texto:
            return p
    raise KeyError(texto)

# ------------------------------------------------------------------- montagem
def main():
    if not os.path.exists(TEMPLATE):
        sys.exit("Template nao encontrado: %s\nPasse o caminho como argumento." % TEMPLATE)

    lista = carregar_secoes()
    secs = {s["titulo"]: s for s in lista}
    d = Document(TEMPLATE)

    # o titulo do trabalho e o primeiro heading de nivel 1 do markdown
    titulo_md = next(s["titulo"] for s in lista if s["nivel"] == 1)
    if "[PREENCHER" in titulo_md:
        print("  ! atencao: o codigo da unidade ainda nao foi preenchido no titulo")
    elif not re.match(r"^#[^#]+# ", titulo_md):
        print("  ! atencao: o manual pede o codigo da unidade entre cerquilhas antes do")
        print("    titulo (ex.: '#12900# AquaSense: ...'). Titulo atual: " + titulo_md[:60])

    # capa e sumario (controles de conteudo do template)
    for sdt in d.element.body.iter(W + "sdt"):
        for t in sdt.iter(W + "t"):
            if t.text and "TÍTULO DO PROJETO" in t.text:
                t.text = t.text.replace("TÍTULO DO PROJETO", titulo_md)

    acha(d, "Título").text = titulo_md

    # 1.1 equipe
    eq = secs["1.1 EQUIPE"]
    bl = blocos(eq["linhas"])
    unidade = next(x for k, x in bl if k == "p" and x.startswith("Unidade"))
    instrutor = next(x for k, x in bl if k == "p" and x.startswith("Instrutor"))
    pu = acha(d, "Unidade Senai:"); pu.text = ""; runs_formatados(pu, unidade)
    pi = acha(d, "Instrutor orientador:"); pi.text = ""; runs_formatados(pi, instrutor)
    tab = [x for k, x in bl if k == "table"][0]
    tbl = d.tables[0]
    for i, linha in enumerate(tab[1:], start=1):
        for j, cel in enumerate(linha):
            c = tbl.rows[i].cells[j]
            c.text = ""
            runs_formatados(c.paragraphs[0], cel)
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(9)
    while len(tbl.rows) > len(tab):
        tbl._tbl.remove(tbl.rows[-1]._tr)

    # secoes de texto (heading do template -> secao do md)
    escreve(d, secs["1.2 PROBLEMA"], acha(d, "PROBLEMA"),
            placeholder=acha(d, "Resuma a demanda solucionada pelo projeto"))
    escreve(d, secs["1.3.1 Área tecnológica da solução"], acha(d, "ÁREA TECNOLÓGICA DA SOLUÇÃO"))
    escreve(d, secs["1.3.2 Justificativa"], acha(d, "JUSTIFICATIVA"))
    escreve(d, secs["1.3.3 Objetivos"], acha(d, "OBJETIVOS"))
    escreve(d, secs["1.3.4 Desenvolvimento"], acha(d, "DESENVOLVIMENTO"))
    escreve(d, secs["1.3.5 Viabilidade técnica"], acha(d, "VIABILIDADE TÉCNICA"))
    escreve(d, secs["1.3.6 Viabilidade econômica"], acha(d, "VIABILIDADE ECONÔMICA"))
    escreve(d, secs["1.3.7 Resultados e conclusão"], acha(d, "RESULTADOS E CONCLUSÃO"))

    # anexos: BOM entra antes do BMG CANVAS do template
    h_canvas = acha(d, "BMG CANVAS")
    h_bom = d.add_paragraph("ORÇAMENTO ESTIMADO (BOM)", style=h_canvas.style)
    h_canvas._p.addprevious(h_bom._p)
    escreve(d, secs["1.4.1 Orçamento Estimado (BOM)"], h_bom)
    escreve(d, secs["1.4.2 BMG Canvas"], h_canvas)
    escreve(d, secs["1.4.3 Situação de Aprendizagem"], acha(d, "SITUAÇÃO DE APRENDIZAGEM"))

    # referencias (recuo zero, sem justificar — padrao NBR 6023)
    ref = secs["REFERÊNCIAS"]
    hr = d.add_paragraph("REFERÊNCIAS", style="Heading 2")
    for tipo, dado in blocos(ref["linhas"]):
        if tipo != "p":
            continue
        p = d.add_paragraph()
        runs_formatados(p, dado)
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    d.save(SAIDA)
    n_pre = sum(1 for p in d.paragraphs if "[PREENCHER" in p.text)
    print("OK ->", SAIDA)
    print("paragrafos:", len(d.paragraphs), "| tabelas:", len(d.tables),
          "| imagens:", len(d.inline_shapes), "| [PREENCHER]:", n_pre)

if __name__ == "__main__":
    main()
