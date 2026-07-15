# -*- coding: utf-8 -*-
"""Confere o TCC_ENTREGA.docx gerado. Uso: python docs/tcc/verificar_tcc.py"""
import io, os, sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
AQUI = os.path.dirname(os.path.abspath(__file__))
d = Document(os.path.join(AQUI, "TCC_ENTREGA.docx"))

# texto de tudo: corpo, tabelas e controles de conteudo (capa)
partes = [p.text for p in d.paragraphs]
for t in d.tables:
    for r in t.rows:
        for c in r.cells:
            partes.append(c.text)
corpo = "\n".join(partes)
capa = "".join(t.text or "" for sdt in list(d.element.body.iter(W + "sdt"))[:1]
               for t in sdt.iter(W + "t"))

# (descricao, trecho, deve_estar_presente, onde)
CHECAGENS = [
    ("titulo aplicado na capa", "AquaSense", True, capa),
    ("marcador da capa substituido", "TÍTULO DO PROJETO", False, capa),
    ("codigo da unidade no titulo", "#HORTO#", True, capa),
    ("instrutor preenchido", "Instrutor orientador: Jairo", True, corpo),
    ("placeholder de instrutor", "PREENCHER: instrutor", False, corpo),
    ("tabela economica antiga", "R$ 300.000", False, corpo),
    ("payback antigo", "payback estimado em cerca de 6 meses", False, corpo.lower()),
    ("calibracao nao realizada", "conectou o transdutor à bancada", False, corpo),
    ("montagem solar nao realizada", "Marcelo montou o circuito", False, corpo),
    ("montagem PVC nao realizada", "Matheus montou fisicamente", False, corpo),
    ("unidade errada (Mariana)", "Mariana", False, corpo),
    ("escopo morto (MQTT)", "MQTT", False, corpo),
    ("escopo morto (Grafana)", "Grafana", False, corpo),
    ("payback correto", "5,6 meses", True, corpo),
    ("OPEX no texto", "39.155", True, corpo),
    ("investimento correto", "260.000", True, corpo),
    ("prototipo declarado com honestidade", "protótipo funcional de bancada", True, corpo),
    ("anexo BMG Canvas", "BMG Canvas", True, corpo),
    ("anexo Situacao de Aprendizagem", "Situação de Aprendizagem", True, corpo),
    ("referencias ABNT", "GEOKON", True, corpo),
]

falhas = 0
for desc, trecho, esperado, onde in CHECAGENS:
    achou = trecho in onde
    ok = achou == esperado
    falhas += 0 if ok else 1
    print("%s %s" % ("  ok  " if ok else "FALHA ", desc))

print()
print("tabelas: %d | imagens: %d | referencias: %d"
      % (len(d.tables), len(d.inline_shapes), corpo.count("Acesso em:")))
pend = [p.text.strip() for p in d.paragraphs
        if any(m in p.text for m in ("[PREENCHER", "[COLAR AQUI", "[ESCREVER AQUI"))]
print("pendencias marcadas em vermelho no docx:")
for p in pend:
    print("   -", p[:70])
print()
print("RESULTADO: %d falha(s)" % falhas)
sys.exit(1 if falhas else 0)
